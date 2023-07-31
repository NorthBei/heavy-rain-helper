
#------------------------使用參數設定--------------------------
debug    =  0;

#------------------------單機測試的檔案--------------------------
global url_isue, url_city
url_isue = "W-C0034-002.xml"  # 依據事件分類
url_city = "W-C0034-001.xml"  # 依據縣市分類

#-----------------------xml檔前綴網網址--------------------------
global base_url, FHYopen_url, url_isue0, FHY_url, FHYLocal;
base_url= "urn:cwb:gov:tw:cwbcommon:0.1" ;
url_isue0 ="http://opendata.cwb.gov.tw/opendataapi?dataid=W-C0033-002&authorizationkey=CWB-CE1AD5F8-F644-4852-BF8F-BF705DE98CAA";
url_city0 ="http://opendata.cwb.gov.tw/opendataapi?dataid=W-C0033-001&authorizationkey=CWB-CE1AD5F8-F644-4852-BF8F-BF705DE98CAA";
FHY_url = "http://fhy.wra.gov.tw/dmchy/wra/WebCIAComponent/EventManager.aspx?groupList=21%2\52c12%252c99&userId=Wramoea&userName=%253f%25ued65%25u865c%253f%25uf4c3%25uec97%253f&userEmail=&userTel=&title=wramoea&dept=0&system=dmchyV2";

#-----------------------(地區列表)-------------------------------
global location_list, RMOffice_list, RWROffice_list, numSortIndex, dirSortIndex


#縣市政府
location_list=['新北市',\
               '臺北市',\
               '基隆市',\
               '桃園市',\
               '新竹縣',\
               '新竹市',\
               '苗栗縣',\
               '臺中市',\
               '彰化縣',\
               '南投縣',\
               '雲林縣',\
               '嘉義縣',\
               '嘉義市',\
               '臺南市',\
               '高雄市',\
               '屏東縣',\
               '宜蘭縣',\
               '花蓮縣',\
               '臺東縣',\
               '連江縣',\
               '金門縣',\
               '澎湖縣']

#河川局
RMOffice_list = {'第一河川局':['宜蘭縣', '新北市','連江縣'], \
                 '第二河川局':['桃園市','新竹縣','新竹市','苗栗縣'], \
                 '第三河川局':['苗栗縣', '臺中市','南投縣'], \
                 '第四河川局':['彰化縣','南投縣'], \
                 '第五河川局':['雲林縣', '嘉義縣','嘉義市','臺南市'], \
                 '第六河川局':['臺南市','高雄市'], \
                 '第七河川局':['屏東縣', '澎湖縣'], \
                 '第八河川局':['臺東縣','金門縣','連江縣'], \
                 '第九河川局':['花蓮縣'], \
                 '第十河川局':['臺北市','新北市','基隆市'],}

#水資源局
RWROffice_list = {'北區水資源局':['新北市','基隆市','臺北市','宜蘭縣','桃園市','新竹市','新竹縣','花蓮縣','連江縣'], \
                  '中區水資源局':['苗栗縣','臺中市','彰化縣','南投縣','雲林縣','金門縣'], \
                  '南區水資源局':['嘉義縣','嘉義市','臺南市','高雄市','屏東縣','臺東縣','澎湖縣'],}

#北水特
TPSpect_list=['新北市',\
              '臺北市']

numSortIndex = '一二三四五六七八九十'
dirSortIndex = '北中南'
#========================================================


#--------------------------定義值班類型-------------------------
print('選擇值班類型:')
print('一般豪大雨值班請按1')
print('三級應變小組請按2')


# emerg=1/0    三級開/不開

x=0;
while x<1:
    type1 = (input('選擇值班類型:'))
    if (type1 == '2'):
        emerg = 1
        break
    elif (type1 =='1'):
        emerg = 0
        break
    else:
        print ("不是1就是2，你是在輸入三小?!")


#========================================================
            
def EndText(SMS, warningState):

    print('---------------------簡訊內容---------------------')
    print()
    print('簡訊內容:')
    print(SMS)
    print()
    print('---------------------發布地區---------------------')
    #print()
    if (warningState['Rain_typ'] == '外面天氣晴' ) :
        print('外面天氣晴耶耶耶耶耶 !!!')

    elif ('解除' in warningState['Rain_typ']):
        print('事件解除準備閃人嘍~')

    else:
        #print('發送對象:')
        #print()
        Institution = warningState['recipient'][0][-2::]  #一開始的接收地區是 "地方政府" -> 取最後兩個字 "政府"
        for City in warningState['recipient']:
            if (City.find(Institution) == -1 ):  # 當居收的地方為 "河川局" 會找不到 "政府"這兩個字 -> 控格格開這兩種單位
                Institution = City[-2::]
                print()
    
            print(City)
        
    print('------------------按任意鍵退出--------------------')
    #print()
    return         
#========================================================
# 產生SMS

import unicodedata    #全形轉半形

def writingSMS(warningState, OpenDateClock):
    if (OpenDateClock=='') :
        OpenDateClock = '1911/[]/[] 上午 []:00:00'
    
    if ( '解除' not in warningState['Rain_typ'] ) :
        if ( '豪雨' in warningState['Rain_typ'] ) :
            if (emerg == 0):
                OpenCentre = '<防災中心>'
            elif (emerg == 1):
                OpenCentre = '<防災中心>'

        elif ( '大雨' in warningState['Rain_typ'] ):
            OpenCentre = '<防災中心>'

        else :
            #print('writingSMS() - 外面天氣晴')
            return
        
        SMS_Region = '';

        # 2016_0502 SMS improve testment: to avoid the bug that there are more than one "。" in the content 
        warningState['content'] = warningState['content'][0:warningState['content'].find('。')]

        # 2016_0502 To avoid the condition that there is a '亦' or '易' in front of the word '有' and cause the condition [ 日針對臺中至臺南地區"易"發布大雨特報 ]
        warningState['content'] = warningState['content'].replace('易','').replace('亦','')
                   
        section_Num = warningState['content'].count('；') + 1  # the ; character signs separate the warning content in many pieces
        for i in range(0,section_Num) :
            if ( i==0 ):
                loc_a = -1
                while (warningState['content'].find('日',loc_a+1) != -1): #找到最後一個 "日"
                    loc_a = warningState['content'].find('日',loc_a+1)
            else:
                SMS_Region += '、'
                loc_a = warningState['content'].find('；')
            
            loc_b = warningState['content'].find('有',loc_a)
        
            SMS_Region +=  unicodedata.normalize('NFKC',warningState['content'][loc_a+1:loc_b] ).replace('(','').replace(')','')   # 全部轉半形 -> 抓出地區 -> 把括號( ) 消去

        
        if ( SMS_Region.count('及') > 1 ):                  # 只留最後一個'及'，其它都是用 '、' 分隔
            SMS_Region = SMS_Region.replace('及','、')
            SMS_Region = SMS_Region[::-1].replace('、','及',1)[::-1]


        if ('天' in SMS_Region):   # 有時候會是  由於鋒面來襲(23日, 24日)兩"天"基隆、北海岸...
                SMS_Region = SMS_Region[ SMS_Region.find('天')+1 ::]
                
        
        day_a = warningState['time'].find('-', warningState['time'].find('-')+1 ) +1  # '2015-12-16T03:25:00+08:00' 找第二個 '-'

        day_b = day_a + 2
        day = warningState['time'][day_a:day_b]

        if ( '下午' in OpenDateClock):
            OP_hour = str( 12 + int( OpenDateClock[ OpenDateClock.find(':')-2 : OpenDateClock.find(':') ] ))   # '2015/12/15 下午 12:00:00'
            if (int(OP_hour) >= 24 ) :
                OP_hour = str( int(OP_hour) - 12)

        elif ( '上午' in OpenDateClock ):
            OP_hour = OpenDateClock[ OpenDateClock.find(':')-2 : OpenDateClock.find(':') ]   # '2015/12/15 下午 12:00:00'
            
        OP_min  = '[ ]'

        if (emerg==0):
            SMS = '中央氣象局於今(' + day + ')日針對'+SMS_Region+\
                  '發布'+warningState['Rain_typ']+'。'+OpenCentre
        elif (emerg ==1):
            SMS = '中央氣象局於今(' + day + ')日針對'+SMS_Region+\
                  '發布'+warningState['Rain_typ']+'，經研判有開設必要'+'，水利署應變小組'+ OP_hour + '時' + OP_min +'分三級開設。'+OpenCentre            


    elif ( '解除' in warningState['Rain_typ'] ):
        if ( '豪雨' in warningState['Rain_typ'] ) :
            CloseCentre = '<防災中心>'

        elif ( '大雨' in warningState['Rain_typ'] ):
            CloseCentre = '<防災中心>'

        else :
            print('writingSMS() - 外面天氣晴')
            return
        
        time_MDhm = warningState['time'].replace('-','年',1).replace('-','月',2).replace('T','日').replace(':','時',1).replace(':','分',2)
        time_MDhm = time_MDhm[ time_MDhm.find('年') : time_MDhm.find('分')+1]
        time_Y   = str( int(warningState['time'][ 0 : warningState['time'].find('-')]) - 1911 )
        
        SMS = '中央氣象局於' + time_Y + time_MDhm + '，' + Rain_typ +'事件。'+ CloseCentre


    else: # Rain_typ= '未知情況'
        print('未知情況')
        input()
        SMS = '未知情況'
        

    file = open('簡訊內容.txt','w')
    file.write(SMS)
    file.close()
    
    return SMS
#========================================================
#產生通報單

#def writingfax(warningState):
#    Year_ROC   = str( int(warningState['time'][ 0 : warningState['time'].find('-')]) - 1911 )
#    fax_content='因應中央氣象局'+Year_ROC+'年'+warningState['time'][5:7]+'月'\
#                                          +warningState['time'][8:10]+'日'+warningState['time'][11:13]+'時'\
#                                          +'起發布豪雨特報，經濟部水利署災害緊急應變小組三級開設，請各單位成'\
#                                          +'立應變小組(中心)時，通報本小組，如需支援亦請聯繫本小組';
#    file = open('通報單內容.txt','w')
#    file.write(fax_content)
#    file.close

#========================================================
# 產生word檔


from docx import Document
from datetime import datetime
import unicodedata    #全形轉半形
    
def writingDocx(warningState, OpenDate, FaxTime):

    #-----------------(DOCX)---------------------

    if '解除' in warningState['Rain_typ']:
        #print('下雨事件解除 !')
        return

    elif warningState['Rain_typ'] == '外面天氣晴' :
        #print('writingDocx() - 外面天氣晴')
        #print()
         return 

    elif warningState['Rain_typ'] == '大雨特報' :
        DocName = '大雨SourceWord.docx'
        

    elif warningState['Rain_typ'] == '豪雨特報':      # 豪雨以上如果是替代役開設都是開豪雨 (大豪雨 / 超大豪雨)
        if type1 == '2':
            DocName = '豪雨SourceWord_third.docx'
        else:
            if IsueCondtion == '續開 - 地區轉換':
                DocName = '豪雨SourceWordOld.docx'

            else:
                DocName = '豪雨SourceWord.docx'

    
    try: document = Document(DocName)
    except:
        print('缺少檔案 - ' + DocName)
        input()

        #document = Document('豪雨SourceWord.docx')  # FOR TESTING
        
        # time example: 2015-12-14T21:10:00+08:00

    Year_ROC   = str( int(warningState['time'][ 0 : warningState['time'].find('-')]) - 1911 )
    document.paragraphs[1].runs[0].text = '受理單位：'+'、'.join(warningState['recipient'])

    if warningState['Rain_typ'] == '豪雨特報' and IsueCondtion == '開設班':
        document.paragraphs[2].runs[0].text = '副本單位：'+'、'.join(warningState['NoAlrmRcpt'])+'、行政院災害防救辦公室、水利防災中心' ;
      
    else:
        document.paragraphs[2].runs[0].text = '副本單位：'+ '行政院災害防救辦公室、水利防災中心'
          
    document.paragraphs[4].runs[1].text = FaxTime  #傳真時間
    document.paragraphs[5].runs[0].text = '依據中央氣象局'+Year_ROC+'年'+warningState['time'][5:7]+'月'\
                                          +warningState['time'][8:10]+'日'+warningState['time'][11:13]+'時'\
                                          +warningState['time'][14:16]+'分'+'發佈'+warningState['Rain_typ']+'傳真辦理。'
    document.paragraphs[6].runs[0].text = warningState['content']
    # runs: 一個段落，中文英文算不同段落，不同字體 style 算不同段落
    # runs: https://automatetheboringstuff.com/chapter13/
    # 更改方式是整個runs砍掉重打


    # warningState['time']['month']+warningState['time']['day']
    SavName = OpenDate+warningState['Rain_typ']+'傳真通報'+\
              FaxTime[4:].replace('/','月').replace(' ','日').replace(':','時')+'分'
    try:
        document.save(SavName +'.docx')

    except:
        document.save(SavName +'.docx')
        print('檔案儲存發生錯誤，請吃土')
        input()


#========================================================
# 顯示事件情況

def PrintIzueCondition(IsueCondtion):
    print('情況：' + IsueCondtion)
    if ( '地區轉換' in IsueCondtion ) or ( '事件轉換' in IsueCondtion) :
        for i in range(0,3): print('      '+IsueCondtion);
        
    return

#========================================================
# 判別 事件/地區 轉換

from datetime import datetime
from datetime import timedelta

def ListSave(List, ListName):
    file = open(ListName+'.txt','w')
    for elementz in List: file.write("%s\n" % elementz);
    file.close()
    return

def RecipientMerge(RecipientPrev, recipient):

    recipientCombined = RecipientPrev
    for Institute in recipient:
        if ( Institute not in RecipientPrev ): recipientCombined.append(Institute)

    return recipientCombined


def RecipientCompare(RecipientPrev, recipient):

    for Institute in recipient:
        if ( Institute not in RecipientPrev ): return '地區轉換';      
    return '地區不變'


def ListLoad(ListName):
    try:
        file = open(ListName +'.txt','r')
        List = file.read().splitlines()   # 存的時候有存成 [OpenName,OpenDate] 的格式
        file.close()

    except:
        List = 'ListLoadingError'

    return List


def FromTextToTimeData(DateInText):   # DateInText = '2015/12/15 上午 12:00:00'

    if (OpenDataParsed == []):
        OpenDate = '開設網站查無資料'
        OpenName = '';
        
    split_a = DateInText.find('/')
    split_b = DateInText.find('/', split_a+1)
    space_a = DateInText.find(' ')
    space_b = DateInText.find(' ', space_a+1)
    colon_a = DateInText.find(':')
    colon_b = DateInText.find(':', colon_a+1)

    DateInDatetime = datetime( int(DateInText[0:split_a]) , int(DateInText[split_a+1 : split_b]) ,int(DateInText[split_b+1 : space_a]) ,\
                               int(DateInText[space_b+1 : colon_a]) , int(DateInText[colon_a+1 : colon_b]) )             
                               
    return DateInDatetime
                               

def FromParserToOpenDate(OpenDataParsed):      

    if (OpenDataParsed == []):
        OpenDateClock = '開設網站查無資料'
        OpenName = '';

    else:
        for Data_each in OpenDataParsed:          
            if '/' in Data_each and ':' in Data_each:    # 找出 '2015/12/15 下午 12:00:00' 這一項
                OpenDateClock = Data_each

            elif ('雨' in Data_each):                    # 找出 '2015_1215大雨'
                OpenName = Data_each

    if 'OpenName' not in locals():
        OpenName = '大雨豪雨以外的事件'
        
    return OpenDateClock, OpenName



def ConditionVerify(OpenDataParsed, Rain_typ, recipient, MinAfterOpen = 15):

    #-------------------------------------判別當下是否有事件，若沒事件則跳出
    if ( Rain_typ== '外面天氣晴' ) :
        OpenDate = ''
        IsueCondtion = '無事件'
        OpenDateClock = ''
        return OpenDate, IsueCondtion, OpenDateClock
        

    elif ('解除' in Rain_typ ):
        OpenDate = ''
        IsueCondtion = '事件解除'
        OpenDateClock = ''
        return OpenDate, IsueCondtion, OpenDateClock


    elif (OpenDataParsed =='開設網頁解析發生錯誤，請手動輸入開設時間'):
        OpenDate = '[ ][ ][ ][ ]'
        IsueCondtion = '你自己判斷唄'
        OpenDateClock = ''
        return OpenDate, IsueCondtion, OpenDateClock

    #-------------------------------------資料處理

    [OpenDateClock, OpenName] = FromParserToOpenDate(OpenDataParsed)
    
    #-------------------------------------判別是否為開設就值行程式-> 抓不到OpenDate -> 開設班
    
    if ( OpenDateClock == '開設網站查無資料' ):
        # 還未開設(無開設資料)
        MonthStr = str(datetime.now().month) if (len(str(datetime.now().month))==2) else "0"+str(datetime.now().month)
        DayStr   = str(datetime.now().day) if (len(str(datetime.now().day))==2) else "0"+str(datetime.now().day)
        OpenDate = MonthStr + DayStr
        IsueCondtion = '開設班'
        ListSave(recipient , 'RecipientPrev')
        return OpenDate, IsueCondtion, OpenDateClock

        
    #-------------------------------------判別當下是 "剛開設 / 續發 / 事件轉換"

    OpenDatetime = FromTextToTimeData(OpenDateClock)
    
    if ( datetime.now() <=  OpenDatetime + timedelta(minutes = MinAfterOpen) ):
        # 還未開設(無開設資料) or 在開設後的15min以內 -> 開設班
        ListSave(recipient , 'RecipientPrev')
        MonthStr = str(OpenDatetime.month) if (len(str(OpenDatetime.month))==2) else "0"+str(OpenDatetime.month)
        DayStr   = str(OpenDatetime.day) if (len(str(OpenDatetime.day))==2) else "0"+str(OpenDatetime.day)
        OpenDate = MonthStr + DayStr
        IsueCondtion = '開設班'
        return OpenDate, IsueCondtion, OpenDateClock
    
    
    elif ( datetime.now() >  OpenDatetime + timedelta(minutes = MinAfterOpen) ):  # 續開大雨 or 豪雨
        if ( OpenName[ OpenName.find('雨')-1 : OpenName.find('雨')+1 ] in Rain_typ ):    # 看續開的雨類別和原本是否相同 Rain_typ ='大雨特報' or '豪雨特報'

            try:
                RecipientPrev = ListLoad('RecipientPrev');
                if ( RecipientPrev == 'ListLoadingError'): IsueCondtion = '續開 - 載入之前發送機關發生錯誤';
                else:                                      IsueCondtion = '續開 - ' + RecipientCompare(RecipientPrev, recipient);
                
                if ( IsueCondtion == '續開 - 地區轉換') or ( IsueCondtion == '續開 - 載入之前發送機關發生錯誤'):
                    recipientCombined = RecipientMerge(RecipientPrev, recipient);
                    ListSave(recipientCombined , 'RecipientPrev');
                #else:  # IsueCondtion == '地區<=上次儲存的''
                
            except:
                IsueCondtion = '載入之前發送機關發生錯誤'
                ListSave(recipient , 'RecipientPrev')
                print(IsueCondtion)
                print()

            MonthStr = str(OpenDatetime.month) if (len(str(OpenDatetime.month))==2) else "0"+str(OpenDatetime.month)
            DayStr   = str(OpenDatetime.day) if (len(str(OpenDatetime.day))==2) else "0"+str(OpenDatetime.day)
            OpenDate = MonthStr + DayStr
            
            return OpenDate, IsueCondtion, OpenDateClock
        

        else:   # 續開的事件和原本相異 -> 豪大雨轉換
            ListSave(recipient , 'RecipientPrev')
            IsueCondtion = '事件轉換 - '  \
                           + OpenName[ OpenName.find('雨')-1 : OpenName.find('雨')+1 ] + ' -> ' \
                           + Rain_typ[ Rain_typ.find('雨')-1 : Rain_typ.find('雨')+1 ] 

            MonthStr = str(OpenDatetime.month) if (len(str(OpenDatetime.month))==2) else "0"+str(OpenDatetime.month)
            DayStr   = str(OpenDatetime.day) if (len(str(OpenDatetime.day))==2) else "0"+str(OpenDatetime.day)
            OpenDate = MonthStr + DayStr
            return OpenDate, IsueCondtion, OpenDateClock

    else:
        ListSave(recipient , 'RecipientPrev')
        IsueCondtion = '未知的情況，請和程式編譯者聯繫'
        MonthStr = str(datetime.now().month) if (len(str(datetime.now().month))==2) else "0"+str(datetime.now().month)
        DayStr   = str(datetime.now().day) if (len(str(datetime.now().day))==2) else "0"+str(datetime.now().day)
        OpenDate = MonthStr + DayStr
        return OpenDate, IsueCondtion, OpenDateClock



#========================================================
# 讀取開設網頁的資訊

from html.parser import HTMLParser  #HTML 
import urllib.request # 抓開設網頁HTML

class MyHTMLParser(HTMLParser):
    def __init__(self):
        HTMLParser.__init__(self)
        self.TagMatch = None
        self.DataParsed = []
        
    def handle_starttag(self, tag, attributes):
        if tag == 'td' or 'TD':
            self.TagMatch = True
        else:
            self.TagMatch = False

            # HTMLParser內建函數：確認 Start Tag 類型是否符合 <script type="text/javascript">
            # len(attributes)==1 ，由於所需要的降雨變數放在 <script type="text/javascript"> 裡面 
            # 所以不符合單一type 的 script要篩去 e.g. <script type1 ='a'type2='b' type3==c >
            # attributes 格式 = [('type','text/javascript')], 為包含n個 1x2 tuple 的 1xn list (n=type 數目)
    
    def handle_endtag(self, tag):
        if (tag == 'td' or 'TD') and self.TagMatch : 
            self.TagMatch = True
            # print('Start / End Tag Matching' )
        else:
            self.TagMatch = False

             # HTMLParser內建函數：確認 End Tag 類型是否符合 </script>
             # 避免開頭結尾不同網頁語法錯誤
            
          
    def handle_data(self, data):    # HTMLParser內建函數：確認完 Start 和 End Tag 後才資料讀取
        if self.TagMatch:
            data = data.replace('\r','').replace('\n','').replace('\t','')
            #print(data)
            if data[0:4] == str(datetime.now().year) or data[0:4] == str(datetime.now().year-1) :
                #print(data)
                self.DataParsed.append(data)

def FHYParser(url):
    try:
        response =  urllib.request.urlopen(url)
        html_contents = response.read().decode('big5')
        
        parser = MyHTMLParser()
        parser.feed(html_contents)
        OpenDataParsed = parser.DataParsed;
        return OpenDataParsed;
    
    except:
        OpenDataParsed = '開設網頁解析發生錯誤，請手動輸入開設時間';
        print(OpenDataParsed) ;
        return OpenDataParsed

#========================================================
# 選擇適當傳真時間  ex: 現在時間12分 -> 發送時間15分 ; 現在時間13分 -> 發送時間20分


from datetime import datetime
from datetime import timedelta

def FaxTimeEstmr (): 
    NowTime = datetime.now()
    
    if ( NowTime.minute % 5 < 3 ):
        WaitMinute = 5 - (NowTime.minute % 5)
    else:
        WaitMinute = 10 - (NowTime.minute % 5)

    SendTime = NowTime + timedelta(minutes = WaitMinute)
    
    TimeStr = str(SendTime)
    TimeStr = TimeStr.replace('-','/').\
              replace(str(SendTime.year),str(SendTime.year-1911))
    # 把2000-01-12 換成 89/01/12
    
    TimeStr = TimeStr[0:-10] # 去除 second 跟 milliscend

    return TimeStr


#========================================================
#縣市政府判別


def City_ID(All_Location) :


    #-----------------(機構與政府查詢)---------------
    alert_area = [] ;       No_alert_area = []
    alert_government = [] ; No_alert_government = [];
    alert_RMO = [] ;        No_alert_RMO = [];
    alert_RWRO = [] ;       No_alert_RWRO = [];
    alert_Spect = [] ;      No_alert_Spect = []; # 北水特 

       
    # 縣政府警報
    for area in location_list:
        if area in All_Location:
            alert_area.append(area)
            alert_government.append(area+'政府')
            
        else:  # 補集合
            No_alert_area.append(area)
            No_alert_government.append(area+'政府')

                
        
    # 河川局 and 水資源局
    for area in alert_area:
        for RMO in RMOffice_list:
            if(area in RMOffice_list[RMO] and not(RMO in alert_RMO)):
                alert_RMO.append(RMO)
                
        for RWRO in RWROffice_list:
            if(area in RWROffice_list[RWRO] and not(RWRO in alert_RWRO)):
                alert_RWRO.append(RWRO)


    for area in No_alert_area: #補集合
        for RMO in RMOffice_list:
            if(area in RMOffice_list[RMO] and not(RMO in No_alert_RMO)):
                No_alert_RMO.append(RMO)
                
        for RWRO in RWROffice_list:
            if(area in RWROffice_list[RWRO] and not(RWRO in No_alert_RWRO)):
                No_alert_RWRO.append(RWRO)

    No_alert_RMO= list(set(No_alert_RMO).difference(set(alert_RMO)))
    No_alert_RWRO= list(set(No_alert_RWRO).difference(set(alert_RWRO)))


    # 北水特 and 水規所 2/28
    for area in TPSpect_list:
        if area in All_Location:
            alert_Spect = ['臺北水源特定區管理局']
            break

    if (len(alert_Spect)==0): #補集合
        No_alert_Spect = ['臺北水源特定區管理局']
    
  

    #-----------------(依照順序排列)---------------
                
    alert_RMO = sorted(alert_RMO, key = lambda x: numSortIndex.find(x[1]))
    alert_RWRO = sorted(alert_RWRO, key = lambda x: dirSortIndex.find(x[0]))

    No_alert_RMO = sorted(No_alert_RMO, key = lambda x: numSortIndex.find(x[1]))
    No_alert_RWRO = sorted(No_alert_RWRO, key = lambda x: dirSortIndex.find(x[0]))
    

    # 排序：https://wiki.python.org/moin/HowTo/Sorting


    #-------------------(地區輸出)-----------------

    recipient = alert_government+alert_RMO+alert_RWRO+alert_Spect+['水利規劃試驗所'] # 發送單位
    No_recipient = No_alert_government+No_alert_RMO+No_alert_RWRO+No_alert_Spect
    return recipient , No_recipient


#========================================================
# 抓出地區

def Location_ID(root_city) :
    All_Location =[];
    
    for location in root_city.findall(pT+'dataset/'+pT+'location') :
        for hazards in location.findall(pT+'hazardConditions/'+pT+'hazards') :
            if ('大雨' in hazards[0][1].text) or ('豪雨' in hazards[0][1].text) :
                All_Location.append(location[0].text)
                
    #print('22個縣市中總共',len(All_Location),'個縣市發布 大(豪)雨特報')
    #print()
        
    return All_Location


#========================================================
# 判別事件
import unicodedata    #全形轉半形

def Issue_ID(root_isue) :
    for issueInfo in root_isue.findall(pT+'dataset') :
        issue = [issueInfo.find(pT+'datasetInfo/'+pT+'datasetDescription').text , issueInfo.find(pT+'datasetInfo/'+pT+'issueTime').text ]
        
        if '豪雨' in issue[0] or '大雨' in issue[0]:
            issueconts = issueInfo.find(pT+'contents/'+pT+'content/'+pT+'contentText').text
            content = unicodedata.normalize('NFKC',issueconts).replace(",","，").replace(";","；").replace('\n','').replace(' ','')
            issueTime = issue[1] # 2015-12-14T21:10:00+08:00

            
            if ('大雨' in issue[0]):
                Rain_typ = '大雨特報'
                if ('解除' in content ):
                    Rain_typ = '解除本次' + Rain_typ
                    return Rain_typ, issueTime , content


            elif ('豪雨' in issue[0]):
                Rain_typ = '豪雨特報'
                if  ('解除' in content ):
                    Rain_typ = '解除本次' + Rain_typ
                    return Rain_typ, issueTime , content
                
                return Rain_typ, issueTime , content
            

    try:
        return Rain_typ , issueTime , content   

    except:
        Rain_typ= '外面天氣晴' ; issueTime='' ; content=''
        return  Rain_typ , issueTime , content    # 無大雨事件

def debug(sss) :
	print("")

#========================================================
# xml解析
import xml.etree.ElementTree as ET #
import urllib.request 

def CWB_xml_decode(url, base_url) :
    #url_isue ="http://opendata.cwb.gov.tw/opendataapi?dataid=W-C0033-002&authorizationkey=CWB-CE1AD5F8-F644-4852-BF8F-BF705DE98CAA"
    #url_city ="http://opendata.cwb.gov.tw/opendataapi?dataid=W-C0033-001&authorizationkey=CWB-CE1AD5F8-F644-4852-BF8F-BF705DE98CAA"
    
    #http://opendata.cwb.gov.tw/opendataapi?dataid={dataid}&authorizationkey={apikey}
    #{dataid}為各資料集代碼 (參照：資料清單) - W-C0033-002
    #{apikey}為會員帳號對應之驗證碼 - CWB-1234ABCD-78EF-GH90-12XY-IJKL12345678

    with urllib.request.urlopen(url) as response:
        xml_contents = response.read().decode('utf-8') 
        #xml網頁先解析城成文字

    xml_contents = xml_contents.replace(' xmlns="'+base_url+'"','') 
    # 因為 root_tag為 <cwbopendata xmlns="urn:cwb:gov:tw:cwbcommon:0.1">，
    # 導致每個tag的名稱都是「 {urn:cwb:gov:tw:cwbcommon:0.1}名子」的形式
    # 在這邊把 root_tag改成  <cwbopendata>，就可以免除這個問題

    root = ET.fromstring(xml_contents)
    #解析文字形式的xml檔案
    
    return root



def Local_xml_decode(url):
    tree = ET.parse(url)
    root = tree.getroot()

    return root

#========================================================
# 使用方式 (單機 / 實際)

# usage = 1 ('Online') / 0 ('OFFline' )
def Use_type(online) :
    
    global url_isue, url_city
    global base_url, FHYopen_url, url_isue0, FHY_url, FHYLocal
    
    if online == 0 :
        print('模式：OFF-line')
        try:
            root_isue  = Local_xml_decode(url_isue)
            root_city = Local_xml_decode(url_city)
            try: OpenDataParsed = FHYParser(FHYLocal) ;
            except:
                OpenDataParsed = FHYParser(FHY_url);
                print('讀取本地開設網頁發生錯誤，改為讀取線上開設網頁');
            pT = '{'+base_url+'}'
            
        except:
            print("讀取本地資料發生錯誤，請檢查資料路徑與格式是否正確")
            print()
            pT = 'Error'; root_isue =''; root_city = '';

    else:
        print('模式：ON-line')
        try :
            root_isue = CWB_xml_decode(url_isue0, base_url)
            root_city = CWB_xml_decode(url_city0, base_url)
            OpenDataParsed = FHYParser(FHY_url)
            pT ='';

        except:
            print("讀取線上資料發生錯誤，請檢查氣象局Open api是否有更新")
            print()
            pT = 'Error'; root_isue =''; root_city = '';
        


    return pT, root_isue, root_city, OpenDataParsed 



#========================================================
# 主程式


# ----------------模式設定-----------------------
import os
try:
    f = open('Setting.txt','r')
except:
    print('缺少設定檔 - Setting.txt')

    
try:
    ExeCode = ''
    while ExeCode != 'END\n' :
        # print(ExeCode)
        exec(ExeCode)
        ExeCode = str(f.readline())
    f.close()
    
except:
    online = 1
    MinAfterOpen = 15
    print('載入外部設定發生錯誤')
    print('使用內部設定 On-line 擷取資料')


# ---------------主程式運行----------------------

for i in range(0,1):
    # 抓取氣象局網站資料
    [pT, root_isue, root_city, OpenDataParsed] = Use_type(online);
    if (pT=='Error'): break;
    [Rain_typ, issueTime , issueconts ] = Issue_ID(root_isue)
    if (Rain_typ == '外面天氣晴'): print(); print('離開吧孩子，外面天氣晴'); print(); break


    # 地區與發送機關
    try:
        All_Location = Location_ID(root_city)
        [recipient , No_recipient] = City_ID(All_Location)      
        
    except:
        print('地區解析發生錯誤，很遺憾'); print(); break


    # 事件狀態與開設日期
    [OpenDate, IsueCondtion, OpenDateClock] = ConditionVerify(OpenDataParsed, Rain_typ, recipient, MinAfterOpen)

    y=0;
    while y<1:
        if (type1 == '2'):
            IsOpenState = (input('開設按1，續發按2: '))
            if (IsOpenState == '1'):
                IsueCondtion = '開設班'
                os.remove('RecipientPrev.txt')
                ListSave(recipient, 'RecipientPrev')    
                break
            else:
                break
        
        elif (type1 =='1'):
            IsOpenState = (input('開設按1，續發按2: '))
            if (IsOpenState == '1'):
                IsueCondtion = '開設班'
                os.remove('RecipientPrev.txt')
                ListSave(recipient, 'RecipientPrev')    
                break
            else:
                break
    
        else:
            print ("不是1就是2，你是在輸入三小?!")
           
        break
    PrintIzueCondition(IsueCondtion)


    # 設定傳真時間
    FaxTime = FaxTimeEstmr()


    # 產生Word
    warningState = {'Rain_typ': Rain_typ, 'time':issueTime, 'content':issueconts, 'recipient':recipient , 'NoAlrmRcpt':No_recipient}
    writingDocx(warningState, OpenDate, FaxTime)


    # 產生SMS
    SMS = writingSMS(warningState, OpenDateClock)

    # 產生通報單內容
#    fax_content = writingfax(warningState)
    
    # 結束訊息
    EndText(SMS, warningState)
    
input()





